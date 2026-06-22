"""
笔记服务层 —— 包含 CRUD、向量双写、异步自动标签等核心业务逻辑。
"""
import asyncio
import html
import io
import json
import os
import re
import tempfile
import uuid
import zipfile
from datetime import datetime, timedelta, timezone

from langchain_core.messages import HumanMessage
from sqlalchemy import delete, func, select, update
from sqlalchemy.ext.asyncio import AsyncSession

from core.logger_handler import logger
from services.note_index_service import NoteIndexService
from services.sources.registry import get_source_registry
from models.note import Note
from models.review_record import ReviewRecord
from schemas import NoteCreate, NoteResponse, NoteUpdate
from utils.prompt_loader import load_prompt

# 艾宾浩斯间隔重复数组（天）
INTERVALS = [1, 2, 4, 7, 15, 30]
NOTE_IMPORT_ALLOWED_EXTENSIONS = {".md", ".markdown", ".txt", ".docx", ".doc"}
NOTE_IMPORT_MAX_FILE_SIZE = 20 * 1024 * 1024


class NoteImportError(ValueError):
    """Raised when an uploaded note file cannot be imported."""


def _safe_import_title(filename: str | None, fallback_text: str = "") -> str:
    base_name = os.path.basename(filename or "").strip()
    stem = os.path.splitext(base_name)[0].strip()
    title = stem

    if not title:
        for line in fallback_text.splitlines():
            cleaned = re.sub(r"^#+\s*", "", line).strip()
            if cleaned:
                title = cleaned
                break

    title = re.sub(r'[\\/:*?"<>|]+', "_", title).strip(" ._")
    return (title or "导入笔记")[:80]


def _decode_import_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _normalize_import_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def _plain_text_to_note_html(text: str) -> str:
    normalized = _normalize_import_text(text)
    if not normalized:
        return ""

    blocks = [block.strip() for block in re.split(r"\n{2,}", normalized) if block.strip()]
    return "\n".join(
        f"<p>{html.escape(block).replace(chr(10), '<br>')}</p>" for block in blocks
    )


def _markdown_to_note_html(text: str) -> str:
    normalized = _normalize_import_text(text)
    if not normalized:
        return ""

    try:
        import markdown as markdown_lib

        return markdown_lib.markdown(
            normalized,
            extensions=["extra", "sane_lists", "nl2br"],
            output_format="html5",
        )
    except Exception as e:
        logger.warning(f"Markdown 导入转换失败，回退为纯文本 HTML: {e}")
        return _plain_text_to_note_html(normalized)


def _extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document as DocxDocument
    except Exception as e:
        raise NoteImportError("当前环境缺少 Word 解析能力，请安装 python-docx 后重试") from e

    try:
        document = DocxDocument(io.BytesIO(content))
    except Exception as e:
        raise NoteImportError("Word 文件解析失败，请确认文件未损坏且为 .docx 格式") from e

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def _extract_doc_text(content: bytes) -> str:
    try:
        from unstructured.partition.doc import partition_doc
    except Exception as e:
        raise NoteImportError("当前环境暂不支持旧版 .doc 解析，请另存为 .docx 后导入") from e

    temp_path = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
            temp_file.write(content)
            temp_path = temp_file.name

        elements = partition_doc(filename=temp_path)
        return "\n\n".join(str(element).strip() for element in elements if str(element).strip())
    except NoteImportError:
        raise
    except Exception as e:
        raise NoteImportError("旧版 Word 文件解析失败，请另存为 .docx 后导入") from e
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except OSError:
                pass


def build_imported_note_payload(filename: str | None, content: bytes, category: str | None = None) -> NoteCreate:
    extension = os.path.splitext(filename or "")[1].lower()
    if extension not in NOTE_IMPORT_ALLOWED_EXTENSIONS:
        supported = "、".join(sorted(NOTE_IMPORT_ALLOWED_EXTENSIONS))
        raise NoteImportError(f"文件类型不支持，目前支持 {supported}")

    if not content:
        raise NoteImportError("导入文件为空")
    if len(content) > NOTE_IMPORT_MAX_FILE_SIZE:
        raise NoteImportError("导入文件大小不能超过 20MB")

    if extension in {".md", ".markdown"}:
        raw_text = _decode_import_text(content)
        note_html = _markdown_to_note_html(raw_text)
    elif extension == ".txt":
        raw_text = _decode_import_text(content)
        note_html = _plain_text_to_note_html(raw_text)
    elif extension == ".docx":
        raw_text = _extract_docx_text(content)
        note_html = _plain_text_to_note_html(raw_text)
    else:
        raw_text = _extract_doc_text(content)
        note_html = _plain_text_to_note_html(raw_text)

    if not _normalize_import_text(raw_text):
        raise NoteImportError("文件未解析到可导入的文本内容")

    normalized_category = category.strip() if category else None
    return NoteCreate(
        title=_safe_import_title(filename, raw_text),
        content=note_html,
        category=normalized_category or None,
    )


def _strip_note_html(value: str | None) -> str:
    if not value:
        return ""
    text = re.sub(r"<br\s*/?>", "\n", value, flags=re.IGNORECASE)
    text = re.sub(r"</(p|div|li|h[1-6]|tr|table)>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    return html.unescape(text)


def _normalize_search_text(value: str | None) -> str:
    return re.sub(r"\s+", " ", _strip_note_html(value).lower()).strip()


def _compact_search_text(value: str | None) -> str:
    return re.sub(r"\s+", "", _normalize_search_text(value))


def _is_subsequence(needle: str, haystack: str) -> bool:
    if not needle:
        return False

    index = 0
    for char in haystack:
        if char == needle[index]:
            index += 1
            if index == len(needle):
                return True
    return False


def _rank_note_search_match(note: Note, query: str) -> int | None:
    normalized_query = _normalize_search_text(query)
    compact_query = _compact_search_text(query)
    if not compact_query:
        return None

    title = _normalize_search_text(note.title)
    compact_title = _compact_search_text(note.title)
    category = _normalize_search_text(note.category)
    compact_category = _compact_search_text(note.category)
    tags = [_normalize_search_text(tag) for tag in (note.tags or []) if isinstance(tag, str)]
    compact_tags = [_compact_search_text(tag) for tag in tags]
    content = _normalize_search_text(note.content)
    compact_content = _compact_search_text(note.content)

    if compact_title == compact_query:
        return 0
    if compact_category == compact_query or compact_query in compact_tags:
        return 1
    if normalized_query in title or compact_query in compact_title:
        return 2
    if (
        normalized_query in category
        or compact_query in compact_category
        or any(normalized_query in tag or compact_query in compact_tag for tag, compact_tag in zip(tags, compact_tags))
    ):
        return 3
    if normalized_query in content or compact_query in compact_content:
        return 4

    terms = [term for term in normalized_query.split(" ") if term]
    if len(terms) > 1:
        combined_text = " ".join([title, category, *tags, content])
        if all(term in combined_text for term in terms):
            return 5

    if _is_subsequence(compact_query, compact_title):
        return 6
    if _is_subsequence(compact_query, compact_category) or any(_is_subsequence(compact_query, tag) for tag in compact_tags):
        return 7
    if _is_subsequence(compact_query, compact_content):
        return 8

    return None


def _note_search_sort_key(match: tuple[int, Note]) -> tuple[int, int, float, str]:
    rank, note = match
    updated_at = note.updated_at or note.created_at
    if updated_at is None:
        timestamp = 0.0
    else:
        if updated_at.tzinfo is None:
            updated_at = updated_at.replace(tzinfo=timezone.utc)
        timestamp = updated_at.timestamp()

    return (
        rank,
        0 if note.is_pinned else 1,
        -timestamp,
        _normalize_search_text(note.title),
    )


def _get_next_interval(review_count: int) -> int:
    """
    根据回顾次数返回下一次回顾间隔天数。
    超出预定义数组后固定使用 30 天间隔。
    """
    if review_count < len(INTERVALS):
        return INTERVALS[review_count]
    return INTERVALS[-1]


class NoteService:
    """
    笔记服务 —— 单例模式（模块级实例 note_service）。

    职责：
    - 笔记 CRUD（PostgreSQL 存储）
    - 通过 NoteIndexService 异步同步索引
    - 异步自动标签生成（LLM 后台任务）
    """

    def __init__(self, embed_model=None):
        """
        初始化笔记索引服务。
        :param embed_model: 嵌入模型实例（后台初始化完成后传入）
        """
        self.index_service = NoteIndexService(embed_model=embed_model)

    def _doc_to_response(self, note: Note) -> NoteResponse:
        """
        将 SQLAlchemy ORM 对象转换为 Pydantic 响应模型。
        """
        return NoteResponse(
            id=note.id,
            user_id=note.user_id,
            title=note.title,
            content=note.content,
            tags=note.tags if note.tags else None,
            category=note.category,
            is_pinned=note.is_pinned if note.is_pinned else False,
            created_at=str(note.created_at) if note.created_at else None,
            updated_at=str(note.updated_at) if note.updated_at else None,
        )

    async def _upsert_note_vector(self, note_id: str, user_id: str, title: str, content: str) -> None:
        """
        后台写入单篇笔记向量。失败只影响语义检索，不影响笔记本体保存。
        """
        if not content.strip():
            return

        try:
            await self.index_service.upsert_note(note_id, user_id, title, content)
        except Exception as e:
            logger.error(f"笔记索引失败 note_id={note_id}: {e}")

    async def _refresh_note_vector(self, note_id: str, user_id: str, title: str, content: str) -> None:
        """
        后台刷新笔记向量。内容清空时只删除旧向量。
        """
        try:
            await self.index_service.refresh_note(note_id, user_id, title, content)
        except Exception as e:
            logger.error(f"更新笔记索引失败 note_id={note_id}: {e}")

    async def create_note(self, db: AsyncSession, user_id: str, payload: NoteCreate) -> NoteResponse:
        """
        创建笔记：
        1. PostgreSQL 写入笔记（若用户提供了 tags/category 直接写入）
        2. 立即返回笔记 ID
        3. pgvector 向量写入后台执行
        4. 若用户未提供 tags/category，后台异步任务自动生成
        """
        note_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc)
        note = Note(
            id=note_id,
            user_id=user_id,
            title=payload.title,
            content=payload.content,
            tags=payload.tags,
            category=payload.category,
            is_pinned=False,
            created_at=now,
            updated_at=now,
        )
        db.add(note)
        await db.commit()

        asyncio.create_task(self._upsert_note_vector(note_id, user_id, payload.title, payload.content))

        # 若用户已提供 tags/category，跳过自动标签生成
        user_provided_meta = payload.tags is not None or payload.category is not None
        if not user_provided_meta:
            asyncio.create_task(self._auto_tag_and_review(note_id, user_id, payload.content))

        return self._doc_to_response(note)

    async def import_note(
        self,
        db: AsyncSession,
        user_id: str,
        filename: str | None,
        content: bytes,
        category: str | None = None,
    ) -> NoteResponse:
        """
        从上传文件导入笔记内容，并复用标准创建流程写入数据库和向量库。
        """
        payload = build_imported_note_payload(filename, content, category)
        return await self.create_note(db, user_id, payload)

    async def update_note(self, db: AsyncSession, note_id: str, user_id: str, payload: NoteUpdate) -> NoteResponse | None:
        """
        更新笔记：
        1. 更新 PostgreSQL 中的 title/content/category/tags
        2. 如果 content 变更，删除旧向量并写入新向量
        """
        stmt = select(Note).where(Note.id == note_id, Note.user_id == user_id)
        result = await db.execute(stmt)
        note = result.scalar_one_or_none()
        if not note:
            return None

        content_changed = payload.content is not None

        if payload.title is not None:
            note.title = payload.title
        if payload.content is not None:
            note.content = payload.content
        if payload.category is not None:
            note.category = payload.category
        if payload.tags is not None:
            note.tags = payload.tags
        if payload.is_pinned is not None:
            note.is_pinned = payload.is_pinned

        note.updated_at = datetime.now(timezone.utc)
        await db.commit()

        # content 变更时后台同步向量，避免保存被外部嵌入服务阻塞。
        if content_changed:
            asyncio.create_task(self._refresh_note_vector(note_id, user_id, note.title, note.content))

        return self._doc_to_response(note)

    async def delete_note(self, db: AsyncSession, note_id: str, user_id: str) -> bool:
        """
        删除笔记：
        1. 删除 PostgreSQL 中的笔记（review_records 通过 FK CASCADE 自动删除）
        2. 删除 pgvector 中的向量
        """
        stmt = select(Note).where(Note.id == note_id, Note.user_id == user_id)
        result = await db.execute(stmt)
        note = result.scalar_one_or_none()
        if not note:
            return False

        await db.execute(delete(Note).where(Note.id == note_id, Note.user_id == user_id))
        await db.commit()

        # 清理索引
        try:
            await self.index_service.delete_note(note_id, user_id)
        except Exception as e:
            logger.error(f"删除笔记索引失败 note_id={note_id}: {e}")

        return True

    async def get_note(self, db: AsyncSession, note_id: str, user_id: str) -> NoteResponse | None:
        """
        根据笔记 ID 和用户 ID 获取笔记详情。
        """
        stmt = select(Note).where(Note.id == note_id, Note.user_id == user_id)
        result = await db.execute(stmt)
        note = result.scalar_one_or_none()
        if not note:
            return None
        return self._doc_to_response(note)

    async def list_notes(
        self,
        db: AsyncSession,
        user_id: str,
        page: int = 1,
        page_size: int = 20,
        category: str | None = None,
        tag: str | None = None,
        sort_by: str = "updated_at",
    ) -> tuple[list[NoteResponse], int]:
        """
        分页查询笔记列表，支持按分类筛选和排序。tag 筛选为内存过滤。
        """
        conditions = [Note.user_id == user_id]
        if category:
            conditions.append(Note.category == category)

        # 先查总数
        count_stmt = select(func.count(Note.id)).where(*conditions)
        result = await db.execute(count_stmt)
        total = result.scalar() or 0

        sort_column = {
            "updated_at": Note.updated_at,
            "created_at": Note.created_at,
            "title": Note.title,
        }.get(sort_by, Note.updated_at)

        if sort_by == "title":
            order = sort_column.asc()
        else:
            order = sort_column.desc()

        # 分页查询，置顶优先 + 指定排序
        stmt = (
            select(Note)
            .where(*conditions)
            .order_by(Note.is_pinned.desc(), order)
            .offset((page - 1) * page_size)
            .limit(page_size)
        )
        result = await db.execute(stmt)
        notes = result.scalars().all()

        note_list = [self._doc_to_response(n) for n in notes]

        # tag 为 JSON 数组，在 Python 层面过滤
        if tag:
            note_list = [n for n in note_list if n.tags and tag in n.tags]

        return note_list, total

    async def search_notes(self, db: AsyncSession, user_id: str, query: str, top_k: int = 30) -> list[NoteResponse]:
        """
        文本搜索笔记：准确匹配优先，模糊匹配在后。

        排序优先级：
        1. 标题完全匹配
        2. 分类/标签完全匹配
        3. 标题包含关键词
        4. 分类/标签包含关键词
        5. 正文包含关键词
        6. 分词匹配
        7. 字符顺序模糊匹配
        """
        if not _compact_search_text(query):
            return []

        stmt = select(Note).where(Note.user_id == user_id)
        result = await db.execute(stmt)
        matches: list[tuple[int, Note]] = []

        for note in result.scalars().all():
            rank = _rank_note_search_match(note, query)
            if rank is not None:
                matches.append((rank, note))

        matches.sort(key=_note_search_sort_key)
        return [self._doc_to_response(note) for _, note in matches[:top_k]]

    async def get_related_notes(
        self,
        db: AsyncSession,
        note_id: str,
        user_id: str,
        top_k: int = 3,
    ) -> list[dict]:
        """
        获取与当前笔记语义相似的其他笔记和知识库文档。
        检索通过 SourceRegistry 进行，笔记服务不直接依赖知识库实现。
        """
        note = await self.get_note(db, note_id, user_id)
        if not note:
            return []

        try:
            chunks = await get_source_registry().search(db, user_id, note.content, source_type="mixed", top_k=top_k + 2)
        except Exception as e:
            logger.error(f"检索关联来源失败 note_id={note_id}: {e}")
            return []

        related_items = []
        for chunk in chunks:
            if chunk.source_type == "note" and chunk.source_id == note_id:
                continue
            related_items.append({
                "id": chunk.source_id,
                "title": chunk.title,
                "content_preview": chunk.content[:150],
                "content": chunk.content,
                "similarity": round(chunk.score or 0.0, 4),
                "source": "knowledge_base" if chunk.source_type == "knowledge" else "note",
            })
        return related_items[:top_k]

    @staticmethod
    def _extract_json(text: str) -> str:
        """
        从 LLM 输出中提取 JSON 字符串。
        处理以下情况：
        - JSON 被 markdown 代码块包裹（```json ... ```）
        - JSON 前面有文字描述
        - JSON 后面有文字描述
        """
        import re

        # 尝试匹配 markdown 代码块中的 JSON
        match = re.search(r'```(?:json)?\s*\n(.*?)\n\s*```', text, re.DOTALL)
        if match:
            return match.group(1).strip()

        # 尝试从第一个 { 到最后一个 } 提取 JSON
        start = text.find('{')
        end = text.rfind('}')
        if start != -1 and end != -1 and end > start:
            return text[start:end + 1]

        return text

    async def _auto_tag_and_review(self, note_id: str, user_id: str, content: str):
        """
        后台异步任务：LLM 分析笔记内容 → 生成标签和分类 → 更新 PostgreSQL → 创建回顾记录。

        此方法在 create_note 结束后通过 asyncio.create_task 执行，
        不阻塞用户保存响应。标签延迟出现是设计意图。
        """
        try:
            # 加载 prompt 模板并填充笔记内容
            prompt_template = load_prompt("auto_tag_prompt")
            prompt = prompt_template.replace("{content}", content)

            # 惰性导入避免模块级循环依赖
            from core.background_init import init_manager
            from db.db_config import AsyncSessionLocal
            chat_model = init_manager.chat_model

            response = await chat_model.ainvoke([HumanMessage(content=prompt)])
            raw_output = response.content.strip()

            # 提取 JSON：LLM 输出可能包含前言、markdown代码块等
            json_str = self._extract_json(raw_output)

            # 解析 LLM 返回的 JSON
            result = json.loads(json_str)
            tags = result.get("tags", [])
            category = result.get("category", "life")

            logger.info(f"自动标签生成完成 note_id={note_id}, tags={tags}, category={category}")

            # 写入 PostgreSQL
            async with AsyncSessionLocal() as session:
                stmt = (
                    update(Note)
                    .where(Note.id == note_id, Note.user_id == user_id)
                    .values(tags=tags, category=category)
                )
                await session.execute(stmt)

                # 创建回顾记录（首次间隔 1 天）
                now = datetime.now()
                review = ReviewRecord(
                    id=str(uuid.uuid4()),
                    note_id=note_id,
                    user_id=user_id,
                    next_review_at=now + timedelta(days=1),
                    interval_days=1,
                    review_count=0,
                )
                session.add(review)
                await session.commit()

        except json.JSONDecodeError as e:
            logger.error(f"解析 LLM 标签输出失败 note_id={note_id}, raw={raw_output[:200]}, extracted={json_str[:200]}: {e}")
        except Exception as e:
            logger.error(f"自动标签后台任务失败 note_id={note_id}: {e}")

    async def autocomplete(self, context: str) -> dict:
        """
        AI 内联补全 —— 基于光标前上下文，调用 Ollama qwen3.5:0.8b 快速生成续写文本。

        Args:
            context: 光标前的文本上下文（最多 50 字）

        Returns:
            {"completion": "续写文本", "success": true/false}
        """
        try:
            from langchain_core.messages import HumanMessage

            from core.background_init import init_manager
            chat_model = init_manager.chat_model

            prompt_template = load_prompt("autocomplete_prompt")
            prompt = prompt_template.format(context=context[-200:])  # 最多取最后200字
            response = await chat_model.ainvoke([HumanMessage(content=prompt)])
            completion = response.content.strip()

            # 防止回复重复已有内容
            if completion and context.endswith(completion[:10]):
                completion = completion[10:]

            return {"success": True, "completion": completion}
        except Exception as e:
            logger.error(f"内联补全失败: {e}")
            return {"success": False, "completion": ""}

    async def assist_stream(self, content: str, action: str):
        """
        AI 写作辅助 SSE 流式输出 —— 支持续写/缩写/扩写三种模式。

        Args:
            content: 用户选中的文本
            action: 操作类型 (expand / summarize / continue)

        Yields:
            SSE 事件数据（字符串）
        """
        from langchain_core.messages import HumanMessage

        from core.background_init import init_manager
        chat_model = init_manager.chat_model

        prompt_template = load_prompt("write_assistant_prompt")
        prompt = prompt_template.format(content=content, action=action)

        try:
            async for chunk in chat_model.astream([HumanMessage(content=prompt)]):
                if chunk.content:
                    yield f"data: {chunk.content}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            logger.error(f"写作辅助流式输出失败: {e}")
            yield f"data: [ERROR: {str(e)}]\n\n"

    async def get_category_stats(self, db: AsyncSession, user_id: str) -> dict:
        """
        获取用户的笔记分类统计 —— 动态查询所有存在的分类并计数。
        """
        stmt = select(Note.category, func.count(Note.id)).where(
            Note.user_id == user_id,
            Note.category.isnot(None),
        ).group_by(Note.category)
        result = await db.execute(stmt)
        categories = [{"category": cat, "count": count} for cat, count in result]

        count_stmt = select(func.count(Note.id)).where(
            Note.user_id == user_id,
            Note.category.is_(None),
        )
        result = await db.execute(count_stmt)
        uncategorized = result.scalar() or 0

        total_stmt = select(func.count(Note.id)).where(Note.user_id == user_id)
        result = await db.execute(total_stmt)
        total = result.scalar() or 0

        return {
            "total": total,
            "categories": categories,
            "uncategorized": uncategorized,
        }

    async def delete_category(self, db: AsyncSession, user_id: str, category: str) -> int:
        """
        删除某个分类及其下所有笔记。
        返回被删除的笔记数量。
        """
        stmt = select(Note).where(
            Note.user_id == user_id,
            Note.category == category,
        )
        result = await db.execute(stmt)
        notes = result.scalars().all()
        note_ids = [n.id for n in notes]
        if not note_ids:
            return 0

        await db.execute(
            delete(Note).where(Note.user_id == user_id, Note.category == category)
        )
        await db.commit()

        for nid in note_ids:
            try:
                await self.index_service.delete_note(nid, user_id)
            except Exception as e:
                logger.error(f"删除分类笔记索引失败 note_id={nid}: {e}")

        return len(note_ids)

    async def export_note_markdown(self, db: AsyncSession, note_id: str, user_id: str) -> str | None:
        """
        导出单篇笔记为 Markdown 文本。
        包含 frontmatter 格式的元数据（标题、标签、分类、日期）。
        """
        note = await self.get_note(db, note_id, user_id)
        if not note:
            return None

        lines = ["---"]
        lines.append(f"title: {note.title}")
        if note.tags:
            lines.append(f"tags: [{', '.join(note.tags)}]")
        if note.category:
            lines.append(f"category: {note.category}")
        lines.append(f"created_at: {note.created_at}")
        lines.append(f"updated_at: {note.updated_at}")
        lines.append("---")
        lines.append("")
        lines.append(f"# {note.title}")
        lines.append("")
        lines.append(note.content)

        return "\n".join(lines)


    async def batch_delete_notes(self, db: AsyncSession, user_id: str, note_ids: list[str]) -> int:
        """
        批量删除笔记：
        1. PostgreSQL 批量删除（级联 review_records）
        2. pgvector 逐个清理向量
        返回实际删除数量。
        """
        if not note_ids:
            return 0

        stmt = select(Note).where(Note.id.in_(note_ids), Note.user_id == user_id)
        result = await db.execute(stmt)
        existing = result.scalars().all()
        existing_ids = [n.id for n in existing]

        if not existing_ids:
            return 0

        await db.execute(delete(Note).where(Note.id.in_(existing_ids), Note.user_id == user_id))
        await db.commit()

        for nid in existing_ids:
            try:
                await self.index_service.delete_note(nid, user_id)
            except Exception as e:
                logger.error(f"批量删除索引失败 note_id={nid}: {e}")

        return len(existing_ids)

    async def batch_update_category(
        self, db: AsyncSession, user_id: str, note_ids: list[str], category: str
    ) -> int:
        """
        批量更新笔记分类。
        返回实际更新的数量。
        """
        if not note_ids:
            return 0

        stmt = (
            update(Note)
            .where(Note.id.in_(note_ids), Note.user_id == user_id)
            .values(category=category)
        )
        result = await db.execute(stmt)
        await db.commit()
        return result.rowcount

    async def batch_export_zip(self, db: AsyncSession, user_id: str, note_ids: list[str]) -> bytes:
        """
        批量导出笔记为 ZIP 压缩包（内含 .md 文件）。
        """
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            for nid in note_ids:
                note = await self.get_note(db, nid, user_id)
                if not note:
                    continue
                md = await self.export_note_markdown(db, nid, user_id)
                if not md:
                    continue
                safe_title = re.sub(r'[\\/:*?"<>|]', '_', note.title or nid)[:80]
                zf.writestr(f"{safe_title}.md", md.encode("utf-8"))
        buf.seek(0)
        return buf.getvalue()

    async def batch_update_pin(
        self, db: AsyncSession, user_id: str, note_ids: list[str], is_pinned: bool
    ) -> int:
        """
        批量置顶/取消置顶笔记。
        返回实际更新的数量。
        """
        if not note_ids:
            return 0

        stmt = (
            update(Note)
            .where(Note.id.in_(note_ids), Note.user_id == user_id)
            .values(is_pinned=is_pinned)
        )
        result = await db.execute(stmt)
        await db.commit()
        return result.rowcount


_note_service_instance: "NoteService | None" = None


def get_note_service() -> NoteService:
    """依赖注入工厂函数。"""
    global _note_service_instance
    if _note_service_instance is None:
        from core.background_init import init_manager
        _note_service_instance = init_manager.note_service
    return _note_service_instance
